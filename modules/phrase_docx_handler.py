"""
Phrase (Memsource) Bilingual DOCX Handler

This module handles the import and export of Phrase (formerly Memsource) bilingual DOCX files.
Phrase uses a multi-table format with numbered inline tags.

Format Structure:
- Multiple tables (typically 2 content tables + 3 metadata tables)
- Content tables with 7 columns:
  1. Segment ID (locked, gray D9D9D9)
  2. Empty (locked, gray D9D9D9)
  3. Segment number (locked, gray D9D9D9)
  4. Source text with tags (locked, gray D9D9D9)
  5. Target text with tags (EDITABLE, no shading)
  6. Status code (locked, colored: 774306=99/confirmed, 5B37C3=MT, etc.)
  7. Empty (no shading)

Tag System:
- Simple tags: {N} (e.g., {1}, {2})
- Formatting tags: {N>text<N} (e.g., {1>CAUTION<1})
- Empty formatting: {N><N}
- Closing tag variant: <N}
- Special content: {N>� <N} (non-breaking space), {N>on page N<N} (cross-ref)

Critical for re-import:
- Only Column 5 (target text) should be edited
- All other columns must remain unchanged
- Tags must be preserved in the target
- Cell shading/locking must be maintained
"""

import os
import re
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from lxml import etree
from typing import List, Dict, Tuple, Optional
from copy import deepcopy


class PhraseSegment:
    """
    Represents a Phrase segment with tag information.
    """
    def __init__(self, segment_id: str, segment_num: str, source_text: str,
                 target_text: str = "", status_code: str = "",
                 row_index: int = 0, table_index: int = 0):
        self.segment_id = segment_id
        self.segment_num = segment_num
        self.source_text = source_text  # Plain text with tags as text
        self.target_text = target_text
        self.status_code = status_code
        self.row_index = row_index
        self.table_index = table_index

        # Extract tags from source for validation
        self.source_tags = self._extract_tags(source_text)

    def _extract_tags(self, text: str) -> List[str]:
        """Extract all Phrase tag numbers from text."""
        # Match {N}, {N>...<N}, <N}, {N><N}
        pattern = r'\{(\d+)[>}]|<(\d+)\}'
        matches = re.findall(pattern, text)
        # Flatten tuples and remove empty strings
        return [m for group in matches for m in group if m]

    @property
    def plain_source(self) -> str:
        """Get source text without tags for translation."""
        # Remove all Phrase tag patterns
        text = re.sub(r'\{\d+\}', '', self.source_text)  # {N}
        text = re.sub(r'\{\d+>.*?<\d+\}', '', text)  # {N>...<N}
        text = re.sub(r'<\d+\}', '', text)  # <N}
        text = re.sub(r'\{\d+><\d+\}', '', text)  # {N><N}
        return text.strip()

    def __repr__(self):
        return f"PhraseSegment(id={self.segment_id[:20]}..., num={self.segment_num}, status={self.status_code})"


class PhraseDOCXHandler:
    """
    Handler for Phrase (Memsource) bilingual DOCX files.

    This class provides methods to:
    - Load and parse Phrase bilingual DOCX files
    - Extract source segments with tag markers
    - Update target segments with translations (preserving exact structure)
    - Save modified files ready for re-import to Phrase
    """

    # Phrase tag patterns
    TAG_SIMPLE = re.compile(r'\{\d+\}')  # {1}
    TAG_FORMATTED = re.compile(r'\{\d+>.*?<\d+\}')  # {1>text<1}
    TAG_CLOSING = re.compile(r'<\d+\}')  # <1}
    TAG_EMPTY = re.compile(r'\{\d+><\d+\}')  # {1><1}
    TAG_ALL = re.compile(r'\{\d+(?:>.*?<\d+)?\}|<\d+\}')  # All patterns

    def __init__(self):
        self.doc = None
        self.content_tables = []  # List of (table_obj, table_index) tuples
        self.segments: List[PhraseSegment] = []
        self.file_path = None

    def load(self, file_path: str) -> bool:
        """
        Load a Phrase bilingual DOCX file.

        Args:
            file_path: Path to the Phrase bilingual DOCX file

        Returns:
            bool: True if loaded successfully, False otherwise
        """
        try:
            self.file_path = file_path
            self.doc = Document(file_path)

            if len(self.doc.tables) == 0:
                print(f"ERROR: No tables found in {file_path}")
                return False

            # Find content tables (rows of segments). A Phrase content
            # table has 7 or 8 columns and its first data row's first
            # cell contains a segment ID of the form ``<id>:<index>``
            # (e.g. ``SSOMDWjYi5xvD7wq_dc10:0``). Matches the relaxed
            # detection in detect_phrase_docx() — short documents and
            # the newer 8-column export variant are both valid.
            self.content_tables = []
            for idx, table in enumerate(self.doc.tables):
                rows = table.rows
                if len(rows) >= 2 and len(rows[0].cells) in (7, 8):
                    # Check if first cell looks like a Phrase segment ID
                    first_cell = rows[0].cells[0].text.strip()
                    if ':' in first_cell:  # Segment IDs have format "xxx:nnn"
                        self.content_tables.append((table, idx))
                        print(f"Found content table {idx} with {len(rows)} rows, {len(rows[0].cells)} columns")

            if not self.content_tables:
                print(f"ERROR: No Phrase content tables found")
                return False

            print(f"Successfully loaded Phrase bilingual DOCX: {file_path}")
            print(f"Content tables: {len(self.content_tables)}")
            print(f"Total segments: {sum(len(t[0].rows) for t in self.content_tables)}")

            return True

        except Exception as e:
            print(f"ERROR loading Phrase DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False

    @staticmethod
    def _unique_cells(row):
        """Return row.cells deduplicated by identity.

        Phrase bilingual DOCX files exist in two on-disk shapes:

          * 7-grid-column tables — every visual column is its own grid
            column, so ``row.cells`` returns 7 distinct cell objects.
          * 8-grid-column tables — the source column is a single
            ``<w:tc>`` with ``<w:gridSpan w:val="2"/>`` spanning two
            adjacent grid columns. python-docx reports this as TWO
            entries in ``row.cells`` that are the *same* cell object
            (``cells[3] is cells[4]``). Word renders 7 visual columns.

        Both shapes carry the same 7 logical columns:
            [ID, ICU, #, Source, Target, Status, Comment]

        Deduplicating by identity collapses the merged duplicate so the
        same fixed indices work for both shapes. Without this, the
        parser silently shifted Target → Status → Comment by one
        column on every 8-grid-column file and corrupted the user's
        translation on export.
        """
        seen = set()
        unique = []
        for cell in row.cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                unique.append(cell)
        return unique

    def extract_source_segments(self) -> List[PhraseSegment]:
        """
        Extract all source segments from the Phrase bilingual DOCX.

        Returns:
            list: List of PhraseSegment objects
        """
        self.segments = []

        if not self.content_tables:
            print("ERROR: No content tables loaded")
            return []

        # Process each content table
        for table_obj, table_idx in self.content_tables:
            for row_idx, row in enumerate(table_obj.rows):
                try:
                    cells = self._unique_cells(row)

                    # Logical columns: [ID, ICU, #, Source, Target, Status, Comment]
                    segment_id = cells[0].text.strip()
                    # Column 1 (ICU) is empty in our use
                    segment_num = cells[2].text.strip()

                    # Extract source and target with formatting as HTML tags
                    source_cell = cells[3]
                    target_cell = cells[4]
                    source_text = self._cell_to_tagged_text(source_cell)
                    target_text = self._cell_to_tagged_text(target_cell)

                    status_code = cells[5].text.strip()
                    # Column 6 (Comment) is read-only — not consumed here

                    # Create PhraseSegment
                    segment = PhraseSegment(
                        segment_id=segment_id,
                        segment_num=segment_num,
                        source_text=source_text,
                        target_text=target_text,
                        status_code=status_code,
                        row_index=row_idx,
                        table_index=table_idx
                    )

                    self.segments.append(segment)

                except Exception as e:
                    print(f"WARNING: Error processing row {row_idx} in table {table_idx}: {e}")
                    continue

        print(f"Extracted {len(self.segments)} segments from Phrase DOCX")
        return self.segments

    def update_target_segments(self, translations: Dict[str, str]) -> int:
        """
        Update target segments with translations.

        Args:
            translations: Dict mapping segment_id to translated text (with Phrase tags)

        Returns:
            int: Number of segments updated
        """
        updated_count = 0

        # Build a lookup map: segment_id -> (table_obj, row_idx)
        # The ID column is always the first logical column — no gridSpan
        # merge there, so row.cells[0] is fine in both 7- and 8-grid
        # variants.
        segment_map = {}
        for table_obj, table_idx in self.content_tables:
            for row_idx, row in enumerate(table_obj.rows):
                segment_id = row.cells[0].text.strip()
                segment_map[segment_id] = (table_obj, row_idx)

        # Update translations. Use _unique_cells to handle the
        # 8-grid-column variant where the Source column has gridSpan=2
        # (Word shows 7 visual columns, python-docx returns 8 with a
        # duplicate entry for Source). Writing into cells[4] in that
        # case would write into the source-duplicate slot, not Target —
        # the bug a user reported on 2026-05-11 with a 63-segment file.
        for segment_id, translation in translations.items():
            if segment_id in segment_map:
                table_obj, row_idx = segment_map[segment_id]
                row = table_obj.rows[row_idx]
                cells = self._unique_cells(row)
                source_cell = cells[3]  # Logical Source column
                target_cell = cells[4]  # Logical Target column

                # Clear existing target content
                self._clear_cell(target_cell)

                # Write new translation copying formatting from source
                self._set_cell_text_with_source_formatting(target_cell, translation, source_cell)

                updated_count += 1

        print(f"Updated {updated_count} target segments")
        return updated_count

    def _clear_cell(self, cell):
        """Clear all content from a cell."""
        for para in cell.paragraphs:
            for run in list(para.runs):
                run._r.getparent().remove(run._r)

    def _set_cell_text(self, cell, text: str):
        """Set cell text, preserving whitespace."""
        if not cell.paragraphs:
            return

        para = cell.paragraphs[0]

        # Clear existing runs
        for run in list(para.runs):
            run._r.getparent().remove(run._r)

        # Add new text with xml:space='preserve' for proper whitespace handling
        if text:
            run = para.add_run(text)
            t_elem = run._r.find(qn('w:t'))
            if t_elem is not None:
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    def _cell_to_tagged_text(self, cell) -> str:
        """
        Convert cell with formatting to HTML-tagged text.
        Uses the same format as memoQ handler: <b>, <i>, <u> tags.
        """
        result_parts = []

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue

                # Determine which tags to apply
                is_bold = run.bold == True
                is_italic = run.italic == True
                is_underline = run.underline == True

                # Build tagged text
                if is_bold or is_italic or is_underline:
                    # Open tags (order: bold, italic, underline)
                    if is_bold:
                        text = f"<b>{text}"
                    if is_italic:
                        text = f"<i>{text}" if not is_bold else text.replace("<b>", "<b><i>", 1)
                    if is_underline:
                        if is_bold and is_italic:
                            text = text.replace("<b><i>", "<b><i><u>", 1)
                        elif is_bold:
                            text = text.replace("<b>", "<b><u>", 1)
                        elif is_italic:
                            text = text.replace("<i>", "<i><u>", 1)
                        else:
                            text = f"<u>{text}"

                    # Close tags (reverse order: underline, italic, bold)
                    if is_underline:
                        text = f"{text}</u>"
                    if is_italic:
                        text = f"{text}</i>"
                    if is_bold:
                        text = f"{text}</b>"

                result_parts.append(text)

        return ''.join(result_parts)

    def _tagged_text_to_runs(self, text: str) -> list:
        """
        Parse text with HTML formatting tags and return a list of runs with formatting info.
        Compatible with Supervertaler's memoQ format.
        """
        import re

        runs = []

        # Track current formatting state
        is_bold = False
        is_italic = False
        is_underline = False

        # Pattern to match opening/closing tags
        tag_pattern = re.compile(r'(</?[biu]>)')

        # Split text by tags, keeping the tags as delimiters
        parts = tag_pattern.split(text)

        current_text = ""

        for part in parts:
            if part == "<b>":
                # Save current run if any
                if current_text:
                    runs.append({
                        'text': current_text,
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline
                    })
                    current_text = ""
                is_bold = True
            elif part == "</b>":
                # Save current run if any
                if current_text:
                    runs.append({
                        'text': current_text,
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline
                    })
                    current_text = ""
                is_bold = False
            elif part == "<i>":
                if current_text:
                    runs.append({
                        'text': current_text,
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline
                    })
                    current_text = ""
                is_italic = True
            elif part == "</i>":
                if current_text:
                    runs.append({
                        'text': current_text,
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline
                    })
                    current_text = ""
                is_italic = False
            elif part == "<u>":
                if current_text:
                    runs.append({
                        'text': current_text,
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline
                    })
                    current_text = ""
                is_underline = True
            elif part == "</u>":
                if current_text:
                    runs.append({
                        'text': current_text,
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline
                    })
                    current_text = ""
                is_underline = False
            else:
                # Regular text
                current_text += part

        # Don't forget the last run
        if current_text:
            runs.append({
                'text': current_text,
                'bold': is_bold,
                'italic': is_italic,
                'underline': is_underline
            })

        return runs

    def _set_cell_text_with_source_formatting(self, target_cell, text: str, source_cell):
        """
        Set cell text parsing HTML formatting tags.
        This preserves word-level bold, italic, and underline formatting.
        """
        if not target_cell.paragraphs:
            return

        para = target_cell.paragraphs[0]

        # Clear existing runs
        for run in list(para.runs):
            run._r.getparent().remove(run._r)

        # Parse HTML tags and create runs
        runs = self._tagged_text_to_runs(text)

        for run_info in runs:
            run_text = run_info.get('text', '')
            if not run_text:
                continue

            run = para.add_run(run_text)

            # Set xml:space='preserve'
            t_elem = run._r.find(qn('w:t'))
            if t_elem is not None:
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

            # Apply formatting
            if run_info.get('bold'):
                run.bold = True
            if run_info.get('italic'):
                run.italic = True
            if run_info.get('underline'):
                run.underline = True

    def save(self, output_path: str = None) -> bool:
        """
        Save the modified document.

        Args:
            output_path: Path to save to (defaults to original path)

        Returns:
            bool: True if saved successfully
        """
        try:
            save_path = output_path or self.file_path
            self.doc.save(save_path)
            print(f"Saved Phrase bilingual DOCX: {save_path}")
            return True
        except Exception as e:
            print(f"ERROR saving Phrase DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False

    def get_segments_for_translation(self) -> List[Tuple[str, str, str]]:
        """
        Get segments that need translation.

        Returns:
            List of (segment_id, source_text, plain_source) tuples
        """
        result = []
        for seg in self.segments:
            # Include all segments (Phrase doesn't have a clear "Not Translated" status)
            # Users can filter based on status_code if needed
            if not seg.target_text or seg.status_code == "MT":
                result.append((seg.segment_id, seg.source_text, seg.plain_source))
        return result


def detect_phrase_docx(file_path: str) -> bool:
    """
    Detect if a DOCX file is a Phrase bilingual file.

    Phrase bilingual DOCX layout (as exported by phrase.com / Memsource):
      * Several intro tables (instructions, header preamble)
      * One or more *content tables* whose rows are translation segments

    A content table has:
      * 7 or 8 columns (Phrase shipped a 7-column variant historically;
        newer exports add an extra trailing column → 8). Earlier versions
        of this detector required *exactly* 7, which rejected current
        Phrase exports.
      * 2+ data rows. The previous threshold of >100 rows was wrong:
        short documents (e.g. one-page certificates) can have far fewer
        segments and were rejected as "not Phrase".
      * First cell of the first data row contains a Phrase segment ID of
        the form ``<base32-ish>:<index>`` (e.g. ``SSOMDWjYi5xvD7wq_dc10:0``).
        The literal ``':' in first_cell`` test stays — it's distinctive
        enough on its own; header rows in other tables either use ``ID``
        or are empty.
    """
    try:
        doc = Document(file_path)

        if not doc.tables:
            return False

        for table in doc.tables:
            if len(table.rows) >= 2 and len(table.rows[0].cells) in (7, 8):
                first_cell = table.rows[0].cells[0].text.strip()
                if ':' in first_cell:
                    return True

        return False

    except Exception as e:
        print(f"Error detecting Phrase DOCX: {e}")
        return False
