"""Smoke test for Phase E run-splitting in DOCX export.

Verifies _split_docx_run_at_offset and _find_or_split_runs_for_range
produce the right shape of paragraph runs when anchoring a Word
comment to a character range that cuts across runs / mid-run.

Run from repo root:
    python .dev/test_docx_run_split.py
"""
from __future__ import annotations
import sys
import tempfile
import os
from pathlib import Path
from copy import deepcopy

repo_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(repo_root))

from docx import Document
from docx.oxml.ns import qn
from docx.text.run import Run as _DocxRun


# Re-implement the two helpers standalone so we don't have to load all
# of Supervertaler.py. They're verbatim copies of the methods.
def _split_docx_run_at_offset(run, offset):
    text = run.text or ''
    if offset <= 0:
        return (None, run)
    if offset >= len(text):
        return (run, None)
    left_text = text[:offset]
    right_text = text[offset:]
    new_element = deepcopy(run._element)
    for t_elem in list(new_element.findall(qn('w:t'))):
        new_element.remove(t_elem)
    new_t = run._element.makeelement(qn('w:t'), {qn('xml:space'): 'preserve'})
    new_t.text = right_text
    new_element.append(new_t)
    run.text = left_text
    run._element.addnext(new_element)
    new_run = _DocxRun(new_element, run._parent)
    return (run, new_run)


def _find_or_split_runs_for_range(para, start, end):
    if start >= end:
        return []
    result: list = []

    def _walk_once():
        runs = list(para.runs)
        offset = 0
        for run in runs:
            run_text = run.text or ''
            run_len = len(run_text)
            run_start = offset
            run_end = offset + run_len
            offset = run_end
            if run_end <= start or run_start >= end:
                continue
            local_start = max(0, start - run_start)
            local_end = min(run_len, end - run_start)
            if local_start > 0:
                _split_docx_run_at_offset(run, local_start)
                return True
            if local_end < run_len:
                _split_docx_run_at_offset(run, local_end)
                return True
            result.append(run)
        return False

    for _ in range(64):
        result = []
        if not _walk_once():
            break
    return result


# ─── Tests ─────────────────────────────────────────────────────────

def _build_para(doc, runs_spec):
    """Build a paragraph from [(text, bold?), ...] specs.
    Returns the new paragraph."""
    para = doc.add_paragraph()
    for text, bold in runs_spec:
        r = para.add_run(text)
        if bold:
            r.bold = True
    return para


def _para_text(para):
    """Return concatenated run text."""
    return ''.join(r.text or '' for r in para.runs)


# Test 1: Anchor that aligns perfectly with an existing run boundary.
# No splitting should happen.
doc = Document()
para = _build_para(doc, [('Hello ', False), ('world', True), ('!', False)])
assert _para_text(para) == 'Hello world!'
# Anchor "world" → chars [6, 11]
runs = _find_or_split_runs_for_range(para, 6, 11)
assert len(runs) == 1
assert runs[0].text == 'world'
assert len(para.runs) == 3, f'Expected 3 runs (no split), got {len(para.runs)}'
assert _para_text(para) == 'Hello world!'
print('TEST 1 (anchor aligns with run boundary - no split): OK')

# Test 2: Anchor cuts mid-run at the start.
# Anchor "rld" → chars [8, 11] cuts into "world".
doc = Document()
para = _build_para(doc, [('Hello world!', False)])
runs = _find_or_split_runs_for_range(para, 8, 11)
assert len(runs) == 1
assert runs[0].text == 'rld'
assert _para_text(para) == 'Hello world!', f'Expected text preserved, got {_para_text(para)!r}'
print('TEST 2 (anchor cuts mid-run at start): OK')

# Test 3: Anchor cuts mid-run at the end.
# Anchor "Hello wo" → chars [0, 8] cuts inside "Hello world!".
doc = Document()
para = _build_para(doc, [('Hello world!', False)])
runs = _find_or_split_runs_for_range(para, 0, 8)
assert len(runs) == 1
assert runs[0].text == 'Hello wo'
assert _para_text(para) == 'Hello world!'
print('TEST 3 (anchor cuts mid-run at end): OK')

# Test 4: Anchor cuts BOTH ends mid-run.
# Anchor "lo wor" → chars [3, 9].
doc = Document()
para = _build_para(doc, [('Hello world!', False)])
runs = _find_or_split_runs_for_range(para, 3, 9)
assert len(runs) == 1
assert runs[0].text == 'lo wor', f'Expected "lo wor", got {runs[0].text!r}'
assert _para_text(para) == 'Hello world!'
print('TEST 4 (anchor cuts both ends mid-run): OK')

# Test 5: Anchor spans multiple runs (no mid-run cuts).
# Runs: ['Hello ', 'world', '!'] (lens 6, 5, 1)
# Anchor "Hello world" → chars [0, 11] — spans first two runs exactly.
doc = Document()
para = _build_para(doc, [('Hello ', False), ('world', True), ('!', False)])
runs = _find_or_split_runs_for_range(para, 0, 11)
assert len(runs) == 2
assert runs[0].text == 'Hello '
assert runs[1].text == 'world'
assert _para_text(para) == 'Hello world!'
print('TEST 5 (anchor spans 2 full runs): OK')

# Test 6: Anchor spans multiple runs with mid-run cuts at both ends.
# Runs: ['Hello ', 'world', '!']
# Anchor "lo wor" → chars [3, 9] — starts inside run0, ends inside run1.
doc = Document()
para = _build_para(doc, [('Hello ', False), ('world', True), ('!', False)])
runs = _find_or_split_runs_for_range(para, 3, 9)
got_texts = [r.text for r in runs]
assert got_texts == ['lo ', 'wor'], f'Expected ["lo ", "wor"], got {got_texts}'
assert _para_text(para) == 'Hello world!'
print('TEST 6 (anchor cuts mid-run AND spans runs): OK')

# Test 7: Formatting preservation - the new (split) run should keep
# the original run's bold flag.
doc = Document()
para = _build_para(doc, [('bold here', True)])
# Split at offset 4 — left "bold", right " here"
_split_docx_run_at_offset(para.runs[0], 4)
assert len(para.runs) == 2
assert para.runs[0].text == 'bold'
assert para.runs[1].text == ' here'
assert para.runs[0].bold is True
assert para.runs[1].bold is True, f'Right half should also be bold, got {para.runs[1].bold!r}'
print('TEST 7 (run split preserves bold formatting): OK')

# Test 8: Edge cases for split offset.
doc = Document()
para = _build_para(doc, [('abc', False)])
run = para.runs[0]
left, right = _split_docx_run_at_offset(run, 0)
assert left is None and right is run
left, right = _split_docx_run_at_offset(run, 3)
assert right is None and left is run
print('TEST 8 (split at boundary edges no-ops correctly): OK')

# Test 9: End-to-end Word comment attachment with anchored range,
# using python-docx's add_comment.
doc = Document()
para = _build_para(doc, [('Hello world!', False)])
anchored_runs = _find_or_split_runs_for_range(para, 6, 11)  # "world"
doc.add_comment(runs=anchored_runs, text='note about "world"',
                 author='Test', initials='T')
tmpdir = tempfile.mkdtemp()
path = os.path.join(tmpdir, 'test.docx')
doc.save(path)
# Reopen and inspect.
doc2 = Document(path)
assert len(list(doc2.comments)) == 1
print('TEST 9 (anchored Word comment saves + reopens with 1 comment): OK')

print()
print('ALL DOCX RUN-SPLIT TESTS PASSED')
